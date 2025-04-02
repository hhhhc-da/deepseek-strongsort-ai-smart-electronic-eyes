from docx import Document
from datetime import datetime
import numpy as np
import pandas as pd
import os
import win32com.client
import uuid
import numpy as np

def create_abnormal_behavior_dict(lnpf, rlpf,
                                  user="审核员 A-103",
                                  template_path=os.path.join('source', 'report.docx'),
                                  output_path=os.path.join('pdfs', 'report-{}-{}-{}.pdf'.format(datetime.now(), str(uuid.uuid4()), str(np.random.normal(1,0.3))).replace(':', '-'))
                                  ):
    '''
    用于解析 DeepSeek 输出数据
    '''
    if len(lnpf['plate'].values) != len(rlpf['reply'].values):
        raise RuntimeError("len(lnpf['plate'].values) != len(rlpf['reply'].values)")
    
    pack = []
    for (plate, reply) in zip(lnpf['plate'].values, rlpf['reply'].values):
        sta_rp = reply.split("*")[0]
        if sta_rp == "有交通违法行为":
            sam = {
                "datetime_report": str(datetime.now()).split()[0],
                "plate": plate,
                # 正则表达式保护生成内容质量
                "report": reply[str(reply).find('*')+1:].replace("没有交通违法行为", "存在交通违法行为"),
                "administrator": user,
                "template_path": template_path,
                "report_output_path": output_path
            }
            pack.append(sam)
    
    return pack

def create_windows_format_report(status_dict:dict={
    "datetime_report": str(datetime.now()).split()[0],
    "plate": "津ABCDEF",
    "report": "（这是一段普通的报告内容）",
    "administrator": "审核员 A-103",
    "template_path": os.path.join('source', 'report.docx'),
    "report_output_path": os.path.join('pdfs', 'report-{}.pdf'.format(datetime.now()).replace(':', '-'))}):
    '''
    用于生成一体化证据链的鉴定书, 格式为 pdf 格式, 本代码仅限 MicroSoft Word, 请注意, 仅支持 Windows 平台使用
    '''
    doc = Document(status_dict["template_path"])
    
    # 替换策略
    replacements = {
        "ph_plate": status_dict['plate'],
        "{{ph_rport}}": status_dict['report'],
        "ph_name": status_dict["administrator"],
        "ph_rp_time": status_dict["datetime_report"]
    }

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))
                        run.bold = True
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        for run in cell.paragraphs[0].runs:
                            if key in run.text:
                                run.text = run.text.replace(key, str(value))
                                run.bold = True
    
    # 保存临时 Word 文件
    temp_docx = os.path.join(".cache", "temp_certificate_{}.docx".format(str(uuid.uuid1())))
    doc.save(temp_docx)
    
    # 使用 Word 应用程序另存为 PDF
    word = win32com.client.Dispatch("Word.Application")
    
    doc_word = word.Documents.Open(os.path.abspath(temp_docx))
    pdf_path = os.path.abspath(status_dict['report_output_path'])
    doc_word.SaveAs(pdf_path, FileFormat=17)
    
    doc_word.Close()
    word.Quit()
    
    os.remove(temp_docx)
    print("PDF 文件已保存: {}\n".format(status_dict['report_output_path']))
    
if __name__ == '__main__':
    # 默认输出即可
    create_windows_format_report()
