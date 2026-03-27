import uuid
import numpy as np
import pandas as pd
from typing import Literal
import os
import time
import requests
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
import sys
import subprocess
import json
import smtplib
import paho.mqtt.client as mqtt
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime
import yaml
import shutil
import traceback

'''
这里的所有内容都可以导出到其他地方作为小插件
这里面的所有内容都不能与其他模块耦合
'''

if sys.platform == "win32":
    import win32com.client
    from docx import Document
else:
    print(f"Warning: 当前系统 ({sys.platform}) 不支持 Word 报告导出功能, 仅限于 Windows 系统")

GLOBAL_YAML_PATH = os.path.join('cfg', 'config.yaml')
GLOBAL_MQTT_TOPIC = "awa"


class DDOSClient():
    '''
    DDOS 客户端, 用于模拟 DDOS 攻击
    主要用于测试系统的抗压能力和攻击检测能力
    '''
    def __init__(self, 
                 target_url="http://127.0.0.1:81/api", 
                 max_threads=50, 
                 request_count=10000,
                 verbose=True
    ):
        self.target_url = target_url
        self.max_threads = max_threads 
        self.request_count = request_count
        
        if verbose:
            self.__report({
                "name": "DDOSClient",
                "target_url": self.target_url,
                "max_threads": self.max_threads,
                "request_count": self.request_count,
                "method": "concurrent.futures.ThreadPoolExecutor"
            })

    def __report(self, args):
        try:
            report_args = {}
            for k, v in args.items():
                if k == 'name':
                    continue
                try:
                    report_args[k] = [v] if not isinstance(v, list) else [str(v[0])]
                except Exception as e:
                    print(f"处理报告项 {k} 失败: {e}")
                    continue

            pf = pd.DataFrame(report_args, index=[args['name']]).T
            print("\n---------------------- DDOSClient 运行报告 ----------------------")
            print(pf)
        except Exception as e:
            print(f"创建报告 DataFrame 失败: {e}")

    @staticmethod
    def send_request(request_id, target_url='http://127.0.0.1:81/api', timeout=10):
        '''
        发送单个HTTP请求并返回结果
        '''
        try:
            start_time = time.time()
            response = requests.get(target_url, timeout=timeout)
            end_time = time.time()
            
            response_time = (end_time - start_time) * 1000
            
            return {
                "request_id": request_id,
                "status": "success",
                "status_code": response.status_code,
                "response_time_ms": round(response_time, 2),
                "thread": threading.current_thread().name
            }
        
        except requests.exceptions.RequestException as e:
            return {
                "request_id": request_id,
                "status": "failed",
                "error": str(e),
                "thread": threading.current_thread().name
            }
        
    def print_stats(self, results):
        '''
        统计请求结果, 主要是用于输出一个结果
        '''
        total = len(results)
        success = sum(1 for r in results if r["status"] == "success")
        failed = total - success
        
        print(f"\n测试统计:")
        print(f"总请求数: {total}")
        print(f"成功请求: {success} ({success/total*100:.1f}%)")
        print(f"失败请求: {failed} ({failed/total*100:.1f}%)")
        
        if success > 0:
            response_times = [r["response_time_ms"] for r in results if r["status"] == "success"]
            print(f"平均响应时间: {sum(response_times)/len(response_times):.2f}ms")
            print(f"最大响应时间: {max(response_times):.2f}ms")
            print(f"最小响应时间: {min(response_times):.2f}ms")

    def attack(self):
        '''
        执行攻击, 不过我们的设备性能可能不足以支持这么高的并发
        '''
        print(f"开始高并发测试 {self.target_url}")
        print(f"配置: 并发线程数={self.max_threads}, 总请求数={self.request_count}")
        print("----------------------------------------")
        
        start_time = time.time()
        results = []
        
        # 使用线程池执行并发请求
        with ThreadPoolExecutor(max_workers=self.max_threads, thread_name_prefix="req-thread") as executor:
            futures = [executor.submit(DDOSClient.send_request, i, self.target_url, 10) for i in range(self.request_count)]
            
            for future in as_completed(futures):
                result = future.result()
                results.append(result)
                
                # if result["status"] == "success":
                #     print(f"请求 #{result['request_id']} [{result['thread']}] 成功, 状态码: {result['status_code']}, 响应时间: {result['response_time_ms']}ms")
                # else:
                #     print(f"请求 #{result['request_id']} [{result['thread']}] 失败: {result['error']}")
        
        total_time = (time.time() - start_time) * 1000
        print("\n----------------------------------------")
        print(f"测试完成，总耗时: {total_time:.2f}ms")
        self.print_stats(results)

class MQTTServer():
    '''
    MQTT 服务器, 用于接收来自监控系统的消息
    主要用于测试系统的消息处理能力和攻击检测能力
    依赖于 mosquitto 子进程, 这里只进行基本的封装
    '''
    def __init__(self, 
                 base_dir=os.path.join('submodules', 'mosquitto'),
                 yaml_path=os.path.join('cfg', 'config.yaml'),
                 verbose=True
    ):
        '''
        配置启动目录, 以及一些基本的 MQTT 配置
        '''
        self.base_dir = base_dir
        self.yaml_path = yaml_path
        self.mqtt_configs = {}

        self.mqtt_client = None

        with open(self.yaml_path, 'r', encoding='utf-8') as f:
            self.configs = yaml.safe_load(f)
            self.mqtt_configs['broker'] = self.configs['mqtt']['broker']
            self.mqtt_configs['port'] = self.configs['mqtt']['port']
            self.mqtt_configs['username'] = self.configs['mqtt']['username']
            self.mqtt_configs['password'] = self.configs['mqtt']['password']

        self.mosquitto_handler = None

        if verbose:
            self.__report({
                "name": "MQTTServer",
                "base_dir": self.base_dir,
                "broker": self.mqtt_configs['broker'],
                "port": self.mqtt_configs['port'],
                "username": self.mqtt_configs['username'],
                "password": "*"*len(self.mqtt_configs['password']) if self.mqtt_configs['password'] else None
            })

    def __report(self, args):
        try:
            report_args = {}
            for k, v in args.items():
                if k == 'name':
                    continue
                try:
                    report_args[k] = [v] if not isinstance(v, list) else [str(v[0])]
                except Exception as e:
                    print(f"处理报告项 {k} 失败: {e}")
                    continue

            pf = pd.DataFrame(report_args, index=[args['name']]).T
            print("\n---------------------- MQTTServer 运行报告 ----------------------")
            print(pf)
        except Exception as e:
            print(f"创建报告 DataFrame 失败: {e}")

    def start_mosquitto(self, config_path=os.path.join('mosquitto.conf')):
        '''
        启动 mosquitto 服务器
        '''
        self.mosquitto_handler = subprocess.Popen(
            ['mosquitto', '-c', config_path], 
            stdout=subprocess.DEVNULL,
            stderr=subprocess.PIPE,
            cwd=self.base_dir
        )

        # 校验进程是否存在
        self.mosquitto_handler.poll()
        if self.mosquitto_handler.returncode is not None:
            stderr = self.mosquitto_handler.stderr.read().decode()
            print(f"无法启动 mosquitto 服务器: {stderr}")
            raise RuntimeError("无法启动 mosquitto 服务器")
        
    def stop_mosquitto(self):
        '''
        停止 mosquitto 服务器
        '''
        if self.mosquitto_handler is not None:
            self.mosquitto_handler.terminate()
            self.mosquitto_handler.wait()

    def publish_message(self, topic, payload):
        '''
        发布 MQTT 消息
        '''
        if self.mqtt_configs is None:
            self.mqtt_client = mqtt.Client()
            self.mqtt_client.username_pw_set(self.mqtt_configs['username'], self.mqtt_configs['password'])
            self.mqtt_client.connect(self.mqtt_configs['broker'], self.mqtt_configs['port'], 60)
            self.mqtt_client.publish(topic, payload)
            self.mqtt_client.disconnect()
        else:
            self.mqtt_client.publish(topic, payload)
            
    def register_mqtt_service(self, on_connect, on_message):
        '''
        注册 MQTT 服务, 所有内容都由外部函数实现
        '''
        self.mqtt_client = mqtt.Client()
        self.mqtt_client.username_pw_set(self.mqtt_configs['username'], self.mqtt_configs['password'])
        self.mqtt_client.on_connect = on_connect
        self.mqtt_client.on_message = on_message

        try:
            self.mqtt_client.connect(self.mqtt_configs['broker'], self.mqtt_configs['port'], 60)
            return True
        except Exception as e:
            print(f"无法连接到 MQTT 服务器: {e}")
            return False

    def stop_mqtt_service(self):
        '''
        停止 MQTT 服务
        '''
        if self.mqtt_client is not None:
            self.mqtt_client.disconnect()
            self.mqtt_client = None

    def fall_in_loop(self):
        '''
        进入 MQTT 循环, 以便于接收消息
        '''
        if self.mqtt_client is not None:
            self.mqtt_client.loop_forever()
        else:
            print("(Warning) 检测到 mqtt_client == None 进入循环失败")

def on_connect(client, userdata, flags, rc):
    global GLOBAL_MQTT_TOPIC
    '''
    这个是连接出结果后的回调函数
    '''
    if isinstance(rc, int) and rc == 0:
        print("MQTT 连接成功")

    if isinstance(GLOBAL_MQTT_TOPIC, str):
        print(f"订阅主题: {GLOBAL_MQTT_TOPIC}")
        client.subscribe(GLOBAL_MQTT_TOPIC)
    else:
        print(f"MQTT 主题配置 GLOBAL_MQTT_TOPIC = {GLOBAL_MQTT_TOPIC} 错误, 无法订阅")

def on_message(client, userdata, msg):
    global GLOBAL_YAML_PATH
    '''
    订阅的主题获得消息之后的回调函数
    '''
    print(f"MQTT 接收到来自 {msg.topic} 主题的消息: {msg.payload.decode()}")

    with open(GLOBAL_YAML_PATH, 'r', encoding='utf-8') as f:
        config = yaml.safe_load(f)

        url = config['http']['server']
        prefix = config['http']['prefix']

    smtp_serve = SMTPClient(yaml_path=GLOBAL_YAML_PATH, subject="安全事件通知", verbose=False)
    
    try:
        data = json.loads(msg.payload.decode())
        # notice = str(data['notice'][0]).decode('utf-8')[7:] if isinstance(data['notice'][0], bytes) else str(data['notice'][0])
        # print("收到的通知:", notice)

        # 我们只需要处理 Event 为 0 的数据
        if int(data['Event']) == 0:
            json_data = {
                "security": '686814B53720631FB57FD83CE57A1D20',
                "code": data['Code']
            }

            # response 形如 {'Code': 0, 'Message': 'Upload successfully', 'Change': 1}
            response = requests.post(url + prefix + '/security', json=json_data)
            response_data = response.json()
            print("接收到的数据:", response_data)
            
            # 如果没有发生改变, 那么我们不需要发送邮箱数据
            if int(response_data['Change']) != 0 and int(data['Code']) != 0:
                timeout = 3
                while timeout > 0:
                    if smtp_serve.send_email(event_type=data['Type'], code=data['Code']):
                        break
                    else:
                        time.sleep(1)
                        timeout -= 1

    except json.JSONDecodeError:
        print("接收到的消息不是有效的 JSON 格式")
    except KeyError:
        print("接收到的消息缺少必要的字段")
    except Exception as e:
        print(f"处理消息时发生错误: {e}")

class SMTPClient():
    '''
    SMTP 客户端, 用于发送邮件通知
    优点在于集中管理, 并不占用太多内存
    值得一提的是, 这个模块必须包裹在 MQTT 消息内
    '''
    def __init__(self, 
                 yaml_path=os.path.join('cfg', 'config.yaml'), 
                 subject="安全事件通知",
                 verbose=True
    ):
        '''
        分开初始化为多个部分, 减少耦合, 以便于后续的动态配置修改
        '''
        self.load_config(yaml_path)
        self.update_subject(subject)

        if verbose:
            self.__report({
                "name": "SMTPClient",
                "yaml_path": self.yaml_path,
                "subject": self.subject,
                "smtp_server": self.smtp_server,
                "smtp_port": self.smtp_port,
                "username": self.username,
                "receiver_email": self.receiver_email
            })

    def __report(self, args):
        try:
            report_args = {}
            for k, v in args.items():
                if k == 'name':
                    continue
                try:
                    report_args[k] = [v] if not isinstance(v, list) else [str(v[0])]
                except Exception as e:
                    print(f"处理报告项 {k} 失败: {e}")
                    continue

            pf = pd.DataFrame(report_args, index=[args['name']]).T
            print("\n---------------------- SMTPClient 运行报告 ----------------------")
            print(pf)
        except Exception as e:
            print(f"创建报告 DataFrame 失败: {e}")

    def update_subject(self, new_subject):
        '''
        更新邮件主题, 一般不会用到
        '''
        self.subject = new_subject

    def load_config(self, yaml_path=None):
        '''
        重新加载配置, 以便于动态修改配置文件
        '''
        if yaml_path is not None:
            self.yaml_path = yaml_path
        else:
            self.yaml_path = os.path.join('cfg', 'config.yaml')

        with open(self.yaml_path, 'r', encoding='utf-8') as f:
            self.config = yaml.safe_load(f)
            self.smtp_server = self.config['smtp']['server']
            self.smtp_port = self.config['smtp']['port']

            self.username = self.config['smtp']['username']
            self.token = self.config['smtp']['token']
            self.receiver_email = self.config['smtp']['receivers']

            # 列表形式
            if isinstance(self.receiver_email, list):
                self.receiver_email = [email.strip() for email in self.receiver_email]
            # 字符串形式
            if isinstance(self.receiver_email, str):
                self.receiver_email = [email.strip() for email in self.receiver_email.split(',')]
            # 空值校验
            if len(self.receiver_email) == 0:
                raise ValueError("接收邮箱列表不能为空")

    def send_email(self, event_type, code) -> bool:
        '''
        发送邮件通知, 需要将列表内的所有用户都发送出去
        '''
        body = "我们检测到了你受到了{}, 认证码为: {}, 时间为: {}.".format(event_type, code, datetime.now())

        message = MIMEMultipart()
        message["From"] = self.username
        message["To"] = ", ".join(self.receiver_email) # 并行发送
        message["Subject"] = self.subject

        message.attach(MIMEText(body, "plain"))

        try:
            with smtplib.SMTP_SSL(self.smtp_server, self.smtp_port) as server:
                server.login(self.username, self.token)
                server.sendmail(self.username, self.receiver_email, message.as_string())
                print("\n邮件发送成功！\n")

                return True
        except Exception as e:
            print(f"\n邮件发送失败: {e}\n")
            return False
    
class ReportExporter():
    '''
    用于导出最后的 Word 报告, 并将其转换为 PDF 格式
    '''
    def __init__(self, 
                 output_dir=os.path.join('runs', 'reports'),
                 verbose=True
    ):
        self.output_dir = output_dir

        # 专门用于管理这个随机数生成器, 以便于在生成报告时使用
        self.certificate_generator = CertificateGnarator(length=100, num_min=0, num_max=80)

        if verbose:
            self.__report({
                "name": "ReportExporter",
                "output_dir": self.output_dir,
                "supported_formats": ['docx', 'pdf'],
                "items": ', '.join(['datetime_report', 'plate', 'report', 'administrator', 'template_path'])
            })

    def __report(self, args):
        try:
            report_args = {}
            for k, v in args.items():
                if k == 'name':
                    continue
                try:
                    report_args[k] = [v] if not isinstance(v, list) else [str(v[0])]
                except Exception as e:
                    print(f"处理报告项 {k} 失败: {e}")
                    continue

            pf = pd.DataFrame(report_args, index=[args['name']]).T
            print("\n---------------------- ReportExporter 运行报告 ----------------------")
            print(pf)
        except Exception as e:
            print(f"创建报告 DataFrame 失败: {e}")

    def export_report(self, 
                      report_name="report", 
                      format:Literal['docx', 'pdf']='pdf',
                      status_dict:dict={
                        "datetime_report": str(datetime.now()).split()[0],
                        "plate": "津ABCDEF",
                        "report": "（这是一段普通的报告内容）",
                        "administrator": "审核员 A-103",
                        "template_path": os.path.abspath(os.path.join('source', 'report.docx')),
                    }):
        '''
        导出 Word 报告, 仅限于 Windows 版本
        所以首先我们需要做系统校验, 然后再考虑使用什么方案
        '''
        if sys.platform != "win32":
            print(f"当前系统 ({sys.platform}) 不支持 Word 报告导出功能, 仅限于 Windows 系统")
            return

        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

        doc = Document(status_dict["template_path"])
        
        # 替换策略 (使用占位符进行替换)
        replacements = {
            "ph_plate": status_dict['plate'],
            "ph_report": status_dict['report'],
            "ph_name": status_dict["administrator"],
            "ph_rp_time": status_dict["datetime_report"]
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))
                            run.bold = True
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            for run in cell.paragraphs[0].runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(value))
                                    run.bold = True
        
        # 保存临时 Word 文件
        temp_docx = os.path.join(".cache", "temp_certificate_{}.docx".format(self.certificate_generator.gen()))
        doc.save(temp_docx)
        # 进程同步需要等一等, 否则会报错
        time.sleep(1)
        
        try:
            if format == 'pdf':
                # 使用 Word 应用程序另存为 PDF
                word = win32com.client.Dispatch("Word.Application")
                
                try:
                    doc_word = word.Documents.Open(os.path.abspath(temp_docx))
                    pdf_path = os.path.join(self.output_dir, f"{report_name}-{datetime.now().strftime('%Y-%m-%d-%H-%M-%S')}.pdf")
                    abs_path = os.path.abspath(pdf_path)
                    doc_word.SaveAs(abs_path, FileFormat=17)
                except Exception as e:
                    print(f"\n转换 PDF 失败: {e}")
                finally:
                    doc_word.Close()
                    word.Quit()
                
                print(f"\nPDF 文件已保存: {pdf_path}\n")
                
                # 等待系统操作
                time.sleep(1)
            else:
                final_docx_path = os.path.join(self.output_dir, f"{report_name}-{datetime.now().strftime('%Y-%m-%d-%H-%M-%S')}.docx")

                shutil.move(temp_docx, final_docx_path)
                print(f"\nWord 文件已保存: {final_docx_path}\n")

                time.sleep(0.3)
        except Exception as e:
            print(f"报告导出失败: {e}")
        finally:
            if os.path.exists(temp_docx):
                os.remove(temp_docx)

class CertificateGnarator():
    '''
    随机序列 ID 生成器, 使用 uuid + numpy 批量生成
    提供缓存功能, 避免出现完全无法调试的情况
    '''
    def __init__(self, length=100, num_min=0, num_max=80):
        '''
        初始化随机数生成器
        '''
        self.length = length
        self.num_min = num_min
        self.num_max = num_max
        
        self.storage = np.random.randint(self.num_min, self.num_max, size=(self.length,))
        self.idx = 0

    def __getitem__(self, idx):
        '''
        获取对应位置数据
        '''
        return self.storage[idx]
    
    def gen(self):
        '''
        生成一个新的随机 ID
        '''
        # 两次检查
        if self.idx > self.length:
            self.storage = np.random.randint(self.num_min, self.num_max, size=(self.length,))
            self.idx = 0
            
        num = self.storage[self.idx]
        self.idx += 1
        
        if self.idx > self.length:
            self.storage = np.random.randint(self.num_min, self.num_max, size=(self.length,))
            self.idx = 0
            
        return str(uuid.uuid3(uuid.NAMESPACE_DNS, str(num)))

    def __call__(self):
        '''
        直接函数调用获取 ID
        '''
        # 两次检查
        if self.idx > self.length:
            self.storage = np.random.randint(self.num_min, self.num_max, size=(self.length,))
            self.idx = 0
            
        num = self.storage[self.idx]
        self.idx += 1
        
        if self.idx > self.length:
            self.storage = np.random.randint(self.num_min, self.num_max, size=(self.length,))
            self.idx = 0
            
        return str(uuid.uuid3(uuid.NAMESPACE_DNS, str(num)))
    
    def get_storage(self):
        '''
        获取当前缓存的随机数列表
        '''
        return self.storage

    def reset(self):
        '''
        重置随机数生成器
        '''
        self.storage = np.random.randint(self.num_min, self.num_max, size=(self.length,))
        self.idx = 0

if __name__ == '__main__':
    mqtt_serve = None

    try:
        mqtt_serve = MQTTServer(
            base_dir = os.path.join('submodules', 'mosquitto'),
            yaml_path = os.path.join('cfg', 'config.yaml'),
            verbose = True
        )
        mqtt_serve.start_mosquitto()
        mqtt_serve.register_mqtt_service(on_connect, on_message)
        
        mqtt_serve.fall_in_loop()
    except KeyboardInterrupt as k:
        print("用户主动停止程序")
    except Exception as e:
        print("接收到:", e)
        traceback.print_exc()
    finally:
        mqtt_serve.stop_mosquitto()
        del mqtt_serve